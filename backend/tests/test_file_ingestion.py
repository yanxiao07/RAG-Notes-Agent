"""真实文件导入测试，覆盖上传、后台入库和格式拒绝路径。"""

from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy.orm import sessionmaker

from app.domain.knowledge.models import Document, DocumentChunk, IngestionJob
from app.extensions.builtin import DocxFileParser


def create_knowledge_base(client: TestClient) -> dict[str, object]:
    response = client.post("/api/v1/knowledge-bases", json={"name": "文件导入测试库"})
    assert response.status_code == 201
    return response.json()


def test_upload_text_file_runs_ingestion_task(client: TestClient) -> None:
    knowledge_base = create_knowledge_base(client)
    response = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("research.md", "# 研究结论\n\n证据必须可追溯。", "text/markdown")},
    )
    assert response.status_code == 202
    job_id = response.json()["ingestionJob"]["id"]

    # TestClient 会等待 BackgroundTasks 完成，随后任务应处于可检索的成功状态。
    job_response = client.get(f"/api/v1/ingestion-jobs/{job_id}")
    assert job_response.status_code == 200
    assert job_response.json()["state"] == "succeeded"

    documents_response = client.get(f"/api/v1/knowledge-bases/{knowledge_base['id']}/documents")
    assert documents_response.status_code == 200
    document = documents_response.json()["items"][0]
    assert document["title"] == "research.md"
    assert document["sourceType"] == "markdown"
    assert document["status"] == "indexed"


def test_upload_persists_selected_enabled_chunker_snapshot(client: TestClient) -> None:
    knowledge_base = create_knowledge_base(client)
    response = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"]), "chunker": "paragraph"},
        files={"file": ("paragraph.txt", "第一段。\n\n第二段。", "text/plain")},
    )

    assert response.status_code == 202
    snapshot = response.json()["ingestionJob"]["configSnapshot"]
    assert snapshot["chunker"] == "paragraph"
    assert snapshot["chunkerVersion"] == "1.0.0"


def test_upload_rejects_unknown_chunker(client: TestClient) -> None:
    knowledge_base = create_knowledge_base(client)
    response = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"]), "chunker": "unknown"},
        files={"file": ("unknown.txt", "内容", "text/plain")},
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "PROCESSING_ERROR"


def test_document_reader_returns_original_markdown_without_leaking_list_payload(
    client: TestClient,
) -> None:
    knowledge_base = create_knowledge_base(client)
    markdown = "# 阅读器验证\n\n```python\nprint('保留代码围栏')\n```"
    created = client.post(
        "/api/v1/documents",
        json={
            "knowledgeBaseId": knowledge_base["id"],
            "title": "reader.md",
            "sourceType": "markdown",
            "content": markdown,
            "parser": "markdown",
        },
    )
    assert created.status_code == 202
    document_id = created.json()["document"]["id"]

    listed = client.get(f"/api/v1/knowledge-bases/{knowledge_base['id']}/documents")
    assert listed.status_code == 200
    assert "rawContent" not in listed.json()["items"][0]

    detail = client.get(f"/api/v1/documents/{document_id}")
    assert detail.status_code == 200
    assert detail.json()["rawContent"] == markdown

    # 不存在的文档统一返回资源不存在，避免详情接口泄露其他工作区资源。
    missing = client.get("/api/v1/documents/not-a-document")
    assert missing.status_code == 404


def test_upload_typora_markdown_keeps_fenced_code_in_raw_document_and_chunks(
    client: TestClient, session_factory: sessionmaker
) -> None:
    knowledge_base = create_knowledge_base(client)
    markdown = (
        "---\ntitle: Typora\n---\n\n# 动态模型\n\n"
        "```python\nfrom langchain_openai import ChatOpenAI\n\n"
        'model = ChatOpenAI(model="gpt-4o-mini", '
        'api_key="sk-abcdefghijklmnopqrstuvwxyz123456")\n```\n'
    )
    response = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("typora.md", markdown, "text/markdown")},
    )
    assert response.status_code == 202
    document_id = response.json()["document"]["id"]

    with session_factory() as session:
        document = session.get(Document, document_id)
        chunks = list(
            session.query(DocumentChunk)
            .filter(DocumentChunk.document_id == document_id)
            .order_by(DocumentChunk.ordinal.asc())
        )
        assert document is not None
        assert "from langchain_openai import ChatOpenAI" in document.raw_content
        assert "sk-abcdefghijklmnopqrstuvwxyz123456" not in document.raw_content
        assert "[REDACTED_SECRET]" in document.raw_content
        assert any("[REDACTED_SECRET]" in chunk.content for chunk in chunks)
        assert any(chunk.metadata_json.get("containsCode") == "true" for chunk in chunks)


def test_upload_rejects_unsupported_file_type(client: TestClient) -> None:
    knowledge_base = create_knowledge_base(client)
    response = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("archive.zip", b"not-a-document", "application/zip")},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "PROCESSING_ERROR"


def test_upload_rejects_duplicate_content_in_the_same_knowledge_base(client: TestClient) -> None:
    knowledge_base = create_knowledge_base(client)
    file_content = "# 可去重资料\n\n同一知识库中相同内容只能入库一次。"
    first = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("first.md", file_content, "text/markdown")},
    )
    assert first.status_code == 202

    duplicate = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("renamed-copy.md", file_content, "text/markdown")},
    )
    assert duplicate.status_code == 409
    assert duplicate.json()["error"]["code"] == "DUPLICATE_RESOURCE"
    assert duplicate.json()["error"]["details"]["documentId"] == first.json()["document"]["id"]


def test_imported_document_can_be_archived_and_reuploaded(client: TestClient) -> None:
    knowledge_base = create_knowledge_base(client)
    file_content = '# 可替换资料\n\n```python\nprint("v1")\n```'
    first = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("first.md", file_content, "text/markdown")},
    )
    assert first.status_code == 202
    document_id = first.json()["document"]["id"]

    archived = client.delete(f"/api/v1/documents/{document_id}")
    assert archived.status_code == 200
    assert archived.json()["status"] == "archived"
    listed = client.get(f"/api/v1/knowledge-bases/{knowledge_base['id']}/documents")
    assert listed.json()["meta"]["total"] == 0

    replacement = client.post(
        "/api/v1/documents/upload",
        data={"knowledge_base_id": str(knowledge_base["id"])},
        files={"file": ("replacement.md", file_content, "text/markdown")},
    )
    assert replacement.status_code == 202
    assert replacement.json()["document"]["id"] != document_id


def test_upload_allows_the_same_content_in_different_knowledge_bases(client: TestClient) -> None:
    first_knowledge_base = create_knowledge_base(client)
    second_knowledge_base = create_knowledge_base(client)
    file_content = "允许跨知识库保留相同资料。"
    for knowledge_base in (first_knowledge_base, second_knowledge_base):
        response = client.post(
            "/api/v1/documents/upload",
            data={"knowledge_base_id": str(knowledge_base["id"])},
            files={"file": ("shared.md", file_content, "text/markdown")},
        )
        assert response.status_code == 202


def test_failed_document_can_be_retried_without_reuploading(
    client: TestClient, session_factory: sessionmaker
) -> None:
    knowledge_base = create_knowledge_base(client)
    created = client.post(
        "/api/v1/documents",
        json={
            "knowledgeBaseId": knowledge_base["id"],
            "title": "retry.md",
            "content": "可重试的入库任务。",
            "sourceType": "markdown",
            "parser": "plain_text",
        },
    )
    assert created.status_code == 202
    document_id = created.json()["document"]["id"]
    with session_factory() as session:
        document = session.get(Document, document_id)
        job = session.query(IngestionJob).filter(IngestionJob.document_id == document_id).one()
        assert document is not None
        document.status = "failed"
        job.state = "failed"
        job.error_code = "INGESTION_FAILED"
        session.commit()

    retried = client.post(f"/api/v1/documents/{document_id}/retry")
    assert retried.status_code == 202

    job_state = client.get(f"/api/v1/ingestion-jobs/{retried.json()['id']}")
    assert job_state.status_code == 200
    assert job_state.json()["state"] == "succeeded"


def test_docx_parser_extracts_paragraphs_and_table_rows() -> None:
    document = DocxDocument()
    document.add_paragraph("研究结论")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "结果"
    stream = BytesIO()
    document.save(stream)

    parsed = DocxFileParser().parse_bytes(title="report.docx", content=stream.getvalue())
    assert "研究结论" in parsed.text
    assert "指标\t结果" in parsed.text
