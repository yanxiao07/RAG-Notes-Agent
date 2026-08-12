"""开发期内置扩展；生产环境可用同契约替换为高级实现。"""

import re
from collections.abc import Iterable
from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.config import get_settings
from app.extensions.contracts import ChunkDraft, ParsedDocument
from app.extensions.registry import ExtensionRegistry
from app.rag.text_normalization import normalize_document_text


class PlainTextParser:
    name = "plain_text"
    version = "1.0.0"

    def parse(self, *, title: str, content: str) -> ParsedDocument:
        # 不可见控制字符会污染向量、导图和 Markdown 渲染，因此在入库边界统一移除。
        normalized = normalize_document_text(content)
        return ParsedDocument(text=normalized, metadata={"title": title})


class Utf8TextFileParser:
    """TXT 和 Markdown 以 UTF-8 读取；显式拒绝损坏编码而非静默替换字符。"""

    name = "utf8_text"
    version = "1.0.0"

    def parse_bytes(self, *, title: str, content: bytes) -> ParsedDocument:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise ValueError("文本文件必须使用 UTF-8 编码。") from exc
        return PlainTextParser().parse(title=title, content=text)


class MarkdownFileParser(Utf8TextFileParser):
    """Typora/标准 Markdown 文件解析器。

    Markdown 在入库边界仍保持原始文本，不尝试转成纯文本；这样围栏代码、
    Front Matter、HTML 图片标签和链接定位信息可以进入证据块并由前端再次渲染。
    """

    name = "markdown"
    version = "1.1.0"

    def parse(self, *, title: str, content: str) -> ParsedDocument:
        parsed = PlainTextParser().parse(title=title, content=content)
        return ParsedDocument(
            text=parsed.text,
            metadata={**parsed.metadata, "format": "markdown"},
        )

    def parse_bytes(self, *, title: str, content: bytes) -> ParsedDocument:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise ValueError("Markdown 文件必须使用 UTF-8 编码。") from exc
        return self.parse(title=title, content=text)


class PdfFileParser:
    name = "pdf"
    version = "1.0.0"

    def parse_bytes(self, *, title: str, content: bytes) -> ParsedDocument:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text(extraction_mode="layout") or "" for page in reader.pages]
        # 页码标记为切分器和引用定位保留轻量结构，不把 PDF 版式控制符带入后续链路。
        text = "\n\n".join(
            f"## 第 {index} 页\n{page.strip()}"
            for index, page in enumerate(pages, start=1)
            if page.strip()
        )
        return ParsedDocument(
            text=normalize_document_text(text),
            metadata={"title": title, "pageCount": str(len(reader.pages))},
        )


class DocxFileParser:
    name = "docx"
    version = "1.0.0"

    def parse_bytes(self, *, title: str, content: bytes) -> ParsedDocument:
        document = DocxDocument(BytesIO(content))
        paragraphs = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        # 表格也经常承载报告中的关键事实，按行保留以便后续切块与引用。
        tables = [
            "\n".join("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            for table in document.tables
        ]
        text = "\n\n".join([*paragraphs, *(table for table in tables if table.strip())])
        return ParsedDocument(text=text, metadata={"title": title})


class ParagraphChunker:
    """按段落优先、按长度兜底的确定性切块器，适合可预测的本地开发和测试。"""

    name = "paragraph"
    version = "1.0.0"

    def __init__(self, max_characters: int = 800) -> None:
        self.max_characters = max_characters

    def chunk(self, document: ParsedDocument) -> list[ChunkDraft]:
        paragraphs = [part.strip() for part in document.text.split("\n\n") if part.strip()]
        chunks: list[ChunkDraft] = []
        buffer = ""
        for paragraph in paragraphs:
            if buffer and len(buffer) + len(paragraph) + 2 > self.max_characters:
                chunks.append(ChunkDraft(ordinal=len(chunks), content=buffer))
                buffer = ""
            # 超长段落采用定长分段，但不会悄悄丢弃任何源文本。
            while len(paragraph) > self.max_characters:
                chunks.append(
                    ChunkDraft(ordinal=len(chunks), content=paragraph[: self.max_characters])
                )
                paragraph = paragraph[self.max_characters :]
            buffer = f"{buffer}\n\n{paragraph}".strip()
        if buffer:
            chunks.append(ChunkDraft(ordinal=len(chunks), content=buffer))
        return chunks


class StructuredChunker:
    """按标题、段落与句子边界切分，并保留有限重叠以降低跨块语义断裂。"""

    name = "structured"
    version = "1.1.0"

    def __init__(self, *, max_characters: int = 900, overlap_characters: int = 120) -> None:
        self.max_characters = max_characters
        self.overlap_characters = overlap_characters

    def chunk(self, document: ParsedDocument) -> list[ChunkDraft]:
        sections = list(self._sections(normalize_document_text(document.text)))
        chunks: list[ChunkDraft] = []
        buffer = ""
        heading = ""
        for section_heading, paragraph in sections:
            # 标题变更意味着语义上下文切换，不能为了凑长度把两个章节混入同一证据块。
            if buffer and section_heading != heading:
                self._append_chunk(chunks, buffer, heading)
                buffer = ""
            unit = self._with_heading(section_heading, paragraph)
            if buffer and len(buffer) + len(unit) + 2 > self.max_characters:
                self._append_chunk(chunks, buffer, heading)
                buffer = "" if self._contains_code_fence(buffer) else self._overlap(buffer)
            for part in self._split_long_unit(unit):
                if buffer and len(buffer) + len(part) + 2 > self.max_characters:
                    self._append_chunk(chunks, buffer, heading)
                    buffer = (
                        ""
                        if self._contains_code_fence(buffer) or self._contains_code_fence(part)
                        else self._overlap(buffer)
                    )
                buffer = f"{buffer}\n\n{part}".strip()
                heading = section_heading or heading
        if buffer:
            self._append_chunk(chunks, buffer, heading)
        return chunks

    @staticmethod
    def _sections(text: str) -> Iterable[tuple[str, str]]:
        heading = ""
        for paragraph in StructuredChunker._markdown_blocks(text):
            if not paragraph:
                continue
            first_line, _, remainder = paragraph.partition("\n")
            if StructuredChunker._is_heading(first_line):
                heading = first_line.lstrip("#").strip()
                if remainder.strip():
                    yield heading, remainder.strip()
                continue
            yield heading, paragraph

    @staticmethod
    def _markdown_blocks(text: str) -> Iterable[str]:
        """仅在围栏外按空行切块，避免 Typora 代码中的空行触发错误分段。"""

        buffer: list[str] = []
        fence: str | None = None
        for line in text.split("\n"):
            stripped = line.strip()
            fence_match = re.match(r"^(`{3,}|~{3,})", stripped)
            if fence is None and fence_match:
                fence = fence_match.group(1)
                buffer.append(line)
                continue
            if fence is not None:
                buffer.append(line)
                if stripped.startswith(fence):
                    fence = None
                continue
            if not stripped:
                if buffer:
                    yield "\n".join(buffer).strip()
                    buffer = []
                continue
            buffer.append(line)
        if buffer:
            yield "\n".join(buffer).strip()

    @staticmethod
    def _is_heading(line: str) -> bool:
        normalized = line.strip()
        return bool(
            normalized.startswith("#")
            or re.match(
                r"^(?:第[一二三四五六七八九十\d]+[章节部分]|"
                r"[一二三四五六七八九十]+、|\d+(?:\.\d+)*[.、]?)\s*\S+",
                normalized,
            )
        )

    @staticmethod
    def _with_heading(heading: str, paragraph: str) -> str:
        return f"{heading}\n{paragraph}".strip() if heading else paragraph

    def _split_long_unit(self, unit: str) -> list[str]:
        if len(unit) <= self.max_characters:
            return [unit]
        if self._contains_code_fence(unit):
            # 代码块优先保持完整；过长时按代码行拆分，并为每个片段补齐围栏，
            # 防止 Markdown 渲染器把后续正文误识别为代码或直接丢弃代码段。
            return self._split_fenced_unit(unit)
        sentences = re.split(r"(?<=[。！？；.!?;])\s+|\n", unit)
        parts: list[str] = []
        buffer = ""
        for sentence in (item.strip() for item in sentences if item.strip()):
            if len(sentence) > self.max_characters:
                if buffer:
                    parts.append(buffer)
                    buffer = ""
                parts.extend(
                    sentence[index : index + self.max_characters]
                    for index in range(0, len(sentence), self.max_characters)
                )
            elif buffer and len(buffer) + len(sentence) + 1 > self.max_characters:
                parts.append(buffer)
                buffer = sentence
            else:
                buffer = f"{buffer} {sentence}".strip()
        if buffer:
            parts.append(buffer)
        return parts or [unit]

    def _append_chunk(self, chunks: list[ChunkDraft], content: str, heading: str) -> None:
        metadata: dict[str, str] = {}
        if heading:
            metadata["section"] = heading
        if self._contains_code_fence(content):
            metadata["containsCode"] = "true"
        chunks.append(
            ChunkDraft(
                ordinal=len(chunks),
                content=content,
                metadata=metadata,
            )
        )

    @staticmethod
    def _contains_code_fence(value: str) -> bool:
        return bool(re.search(r"(?m)^\s*(`{3,}|~{3,})", value))

    def _split_fenced_unit(self, unit: str) -> list[str]:
        lines = unit.splitlines()
        start = next(
            (index for index, line in enumerate(lines) if re.match(r"^\s*(`{3,}|~{3,})", line)),
            None,
        )
        if start is None:
            return [unit]
        opening = lines[start].strip()
        marker_match = re.match(r"(`{3,}|~{3,})", opening)
        if marker_match is None:
            return [unit]
        marker = marker_match.group(1)
        end = next(
            (
                index
                for index in range(start + 1, len(lines))
                if lines[index].strip().startswith(marker)
            ),
            None,
        )
        if end is None:
            # 未闭合围栏也不能丢弃内容，保留原始单元交给上层展示并记录完整文本。
            return [unit]
        prefix = "\n".join(lines[:start]).strip()
        body = lines[start + 1 : end]
        suffix = "\n".join(lines[end + 1 :]).strip()
        capacity = max(self.max_characters - len(opening) - len(marker) - 3, 1)
        groups: list[list[str]] = []
        current: list[str] = []
        current_length = 0
        for line in body:
            line_length = len(line) + (1 if current else 0)
            if current and current_length + line_length > capacity:
                groups.append(current)
                current = []
                current_length = 0
            current.append(line)
            current_length += line_length
        if current or not groups:
            groups.append(current)
        parts: list[str] = []
        for index, group in enumerate(groups):
            code = "\n".join([opening, *group, marker])
            if index == 0 and prefix:
                code = f"{prefix}\n\n{code}"
            if index == len(groups) - 1 and suffix:
                code = f"{code}\n\n{suffix}"
            parts.append(code)
        return parts

    def _overlap(self, content: str) -> str:
        if self.overlap_characters <= 0:
            return ""
        return content[-self.overlap_characters :].lstrip()


def build_builtin_registry(*, enabled_chunkers: set[str] | None = None) -> ExtensionRegistry:
    """注册部署内置扩展，切分器仅按环境白名单开放给入库任务。"""

    active_chunkers = enabled_chunkers or get_settings().enabled_chunker_names
    supported_chunkers = {"paragraph": ParagraphChunker, "structured": StructuredChunker}
    unknown_chunkers = active_chunkers - supported_chunkers.keys()
    if unknown_chunkers:
        raise ValueError(f"未知的内置 Chunker 配置: {','.join(sorted(unknown_chunkers))}")
    if not active_chunkers:
        raise ValueError("至少需要启用一个 Chunker")
    registry = ExtensionRegistry()
    registry.register_parser(PlainTextParser())
    registry.register_parser(MarkdownFileParser())
    registry.register_file_parser(Utf8TextFileParser())
    registry.register_file_parser(MarkdownFileParser())
    registry.register_file_parser(PdfFileParser())
    registry.register_file_parser(DocxFileParser())
    for name in sorted(active_chunkers):
        registry.register_chunker(supported_chunkers[name]())
    return registry
